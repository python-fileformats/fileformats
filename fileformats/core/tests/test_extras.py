import platform
import typing as ty
from pathlib import Path

import pytest
from docx import Document

from fileformats.core import FileSet, MockMixin, extra, extra_implementation
from fileformats.testing import Foo, WithExtra
from fileformats.vendor.openxmlformats_officedocument.application import (
    Wordprocessingml_Document,
)


def test_sample():
    test_inst = Foo.sample()
    assert test_inst.fspath.exists()
    assert test_inst.fspath.suffix == ".foo"


def test_mock():
    mock = Foo.mock()
    if platform.system() == "Windows":
        expected = Path(f"{Path().cwd().drive}\\mock\\foo.foo")
    else:
        expected = Path("/mock/foo.foo")
    assert mock.fspath == expected
    assert not mock.fspath.exists()
    assert isinstance(mock, MockMixin)


class Woo(FileSet):
    @extra
    def test_extra(self, a: int, b: float, c: ty.Optional[str] = None) -> float:
        raise NotImplementedError


def test_extra_signature_no_default():
    extra_implementation(Woo.test_extra)

    def woo_test_extra(woo: Woo, a: int, b: float) -> float:
        pass


def test_extra_signature1():

    with pytest.raises(TypeError, match="missing required argument"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int) -> float:
            pass


def test_extra_signature2():

    with pytest.raises(TypeError, match="name of parameter"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int, d: str) -> float:
            pass


def test_extra_signature3():

    with pytest.raises(TypeError, match="found additional argument"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(
            woo: Woo, a: int, b: float, c: ty.Optional[str], d: int
        ) -> float:
            pass


def test_extra_signature4():

    with pytest.raises(TypeError, match="return type"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(woo: Woo, a: int, b: str) -> int:
            pass


class WooKwargs(FileSet):
    @extra
    def test_extra_kwargs(self, a: int, **kwargs: int) -> float:
        raise NotImplementedError


def test_extra_signature_kwargs_match():
    """A method with **kwargs is matched by an implementation with **kwargs of a
    compatible type."""

    @extra_implementation(WooKwargs.test_extra_kwargs)
    def woo_test_extra_kwargs(woo: WooKwargs, a: int, **kwargs: int) -> float:
        pass


def test_extra_signature_kwargs_impl_missing():
    """An implementation that drops the **kwargs present on the abstract method
    should be rejected."""

    with pytest.raises(TypeError, match="variable keywords vs non-variable keywords"):

        @extra_implementation(WooKwargs.test_extra_kwargs)
        def woo_test_extra_kwargs(woo: WooKwargs, a: int) -> float:
            pass


def test_extra_signature_kwargs_method_missing():
    """An implementation that adds **kwargs not present on the abstract method
    should be rejected."""

    with pytest.raises(TypeError, match="non-variable keywords vs variable keywords"):

        @extra_implementation(Woo.test_extra)
        def woo_test_extra(
            woo: Woo, a: int, b: float, c: ty.Optional[str] = None, **kwargs: ty.Any
        ) -> float:
            pass


def test_extra_signature_kwargs_type_mismatch():
    """**kwargs annotated with incompatible types between the method and the
    implementation should be rejected."""

    with pytest.raises(TypeError, match="Type of keyword args"):

        @extra_implementation(WooKwargs.test_extra_kwargs)
        def woo_test_extra_kwargs(woo: WooKwargs, a: int, **kwargs: str) -> float:
            pass


class WooGeneric(FileSet):
    @extra
    def test_extra_generic(self, a: ty.Mapping[str, int]) -> None:
        raise NotImplementedError


class WooGenericSubtype(WooGeneric):
    pass


class WooGenericPlain(WooGeneric):
    pass


def test_extra_signature_generic_exact_match():
    """A ty.Mapping[str, int]-typed method is matched by an implementation typed
    with the identical generic annotation."""

    @extra_implementation(WooGeneric.test_extra_generic)
    def woo_test_extra_generic(woo: WooGeneric, a: ty.Mapping[str, int]) -> None:
        pass


def test_extra_signature_generic_subtype_origin():
    """A generic subtype of the method's generic origin should be accepted, e.g.
    ty.MutableMapping[str, int] for a method typed as ty.Mapping[str, int]."""

    @extra_implementation(WooGeneric.test_extra_generic)
    def woo_test_extra_generic(
        woo: WooGenericSubtype, a: ty.MutableMapping[str, int]
    ) -> None:
        pass


def test_extra_signature_generic_plain_subclass():
    """A plain (unparameterised) class that is a subclass of the method's generic
    origin should be accepted, e.g. `dict` for a method typed as
    ty.Mapping[str, int]."""

    @extra_implementation(WooGeneric.test_extra_generic)
    def woo_test_extra_generic(woo: WooGenericPlain, a: dict) -> None:
        pass


def test_extra_signature_generic_mismatch():
    """An implementation typed with an unrelated generic should be rejected."""

    with pytest.raises(TypeError, match="Type of 'a' arg"):

        @extra_implementation(WooGeneric.test_extra_generic)
        def woo_test_extra_generic(woo: WooGeneric, a: ty.Sequence[int]) -> None:
            pass


def test_vendor_extra_load(tmp_path: Path):

    fspath = tmp_path / "test.docx"

    # Create a sample document
    data = Document()
    data.add_heading("A document")
    data.add_paragraph("The quick brown fox jumped over the lazy dog.")

    Wordprocessingml_Document.new(fspath, data=data)
    doc = Wordprocessingml_Document(fspath)
    hsh = doc.hash()
    reloaded_data = doc.load()
    assert documents_equal(data, reloaded_data)
    data.add_paragraph("Another paragraph")
    doc.save(data)
    assert doc.hash() != hsh


def documents_equal(doc1: Document, doc2: Document) -> bool:
    """Compare two Document objects by their XML content."""

    # Get the main document part XML
    xml1 = doc1.part.element.xml
    xml2 = doc2.part.element.xml

    return xml1 == xml2


@extra_implementation(WithExtra.foo)
def with_extra_foo_override(wextra: WithExtra, an_arg: int) -> int:
    return an_arg * 3


def test_extra_override():
    """
    Test that extra implementations can be overridden by functions
    outside of the extras module."""
    wextra = WithExtra.sample()
    assert wextra.foo(2) == 6
